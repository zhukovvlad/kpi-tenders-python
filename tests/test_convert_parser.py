"""Unit tests for DOCX and XLSX parsers.

Fixtures are built in-memory using python-docx and openpyxl so there are
no file-system dependencies and the tests run fully offline.
"""

from io import BytesIO

from docx import Document
from openpyxl import Workbook

from app.parsers.docx_parser import count_h1_h2, docx_to_markdown
from app.parsers.xlsx_parser import xlsx_to_markdown_sections

# ---------------------------------------------------------------------------
# Builders
# ---------------------------------------------------------------------------


def make_docx(
    headings: list[tuple[int, str]] | None = None,
    paragraphs: list[str] | None = None,
    tables: list[list[list[str]]] | None = None,
    bold_text: str | None = None,
    italic_text: str | None = None,
    bullet_items: list[str] | None = None,
    numbered_items: list[str] | None = None,
) -> bytes:
    """Build a minimal DOCX file in memory and return its bytes."""
    doc = Document()

    for level, text in headings or []:
        doc.add_heading(text, level=level)

    for text in paragraphs or []:
        doc.add_paragraph(text)

    for rows in tables or []:
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                tbl.rows[i].cells[j].text = str(cell)

    if bold_text is not None:
        run = doc.add_paragraph().add_run(bold_text)
        run.bold = True

    if italic_text is not None:
        run = doc.add_paragraph().add_run(italic_text)
        run.italic = True

    for item in bullet_items or []:
        doc.add_paragraph(item, style="List Bullet")

    for item in numbered_items or []:
        doc.add_paragraph(item, style="List Number")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx(sheets: dict[str, list[list]]) -> bytes:
    """Build a minimal XLSX file in memory and return its bytes."""
    wb = Workbook()
    wb.remove(wb.active)  # remove the default empty sheet
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name)
        for row in rows:
            ws.append(row)
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ===========================================================================
# DOCX parser tests
# ===========================================================================


class TestDocxHeadings:
    def test_heading1(self):
        md = docx_to_markdown(make_docx(headings=[(1, "Первый уровень")]))
        assert "# Первый уровень" in md

    def test_heading2(self):
        md = docx_to_markdown(make_docx(headings=[(2, "Подраздел")]))
        assert "## Подраздел" in md

    def test_heading3(self):
        md = docx_to_markdown(make_docx(headings=[(3, "Пункт")]))
        assert "### Пункт" in md

    def test_headings_preserve_order(self):
        md = docx_to_markdown(make_docx(headings=[(1, "Alpha"), (2, "Beta"), (3, "Gamma")]))
        assert md.index("# Alpha") < md.index("## Beta") < md.index("### Gamma")


class TestDocxParagraphs:
    def test_plain_text(self):
        md = docx_to_markdown(make_docx(paragraphs=["Обычный параграф"]))
        assert "Обычный параграф" in md

    def test_empty_paragraph_skipped(self):
        md = docx_to_markdown(make_docx(paragraphs=["", "  "]))
        # Empty strings should not generate blank placeholder blocks.
        assert md.strip() == ""

    def test_bold_run(self):
        md = docx_to_markdown(make_docx(bold_text="Жирный текст"))
        assert "**Жирный текст**" in md

    def test_italic_run(self):
        md = docx_to_markdown(make_docx(italic_text="Курсив"))
        assert "*Курсив*" in md


class TestDocxLists:
    def test_bullet_list_items(self):
        md = docx_to_markdown(make_docx(bullet_items=["Пункт А", "Пункт Б"]))
        assert "- Пункт А" in md
        assert "- Пункт Б" in md

    def test_numbered_list_items(self):
        md = docx_to_markdown(make_docx(numbered_items=["Первый", "Второй"]))
        assert "1. Первый" in md
        assert "1. Второй" in md


class TestDocxTables:
    def test_table_header_row(self):
        data = [["Имя", "Должность"], ["Иван", "Инженер"]]
        md = docx_to_markdown(make_docx(tables=[data]))
        assert "| Имя | Должность |" in md

    def test_table_separator_row(self):
        data = [["A", "B"], ["1", "2"]]
        md = docx_to_markdown(make_docx(tables=[data]))
        assert "| --- | --- |" in md

    def test_table_data_row(self):
        data = [["Товар", "Цена"], ["Бетон", "5800"]]
        md = docx_to_markdown(make_docx(tables=[data]))
        assert "| Бетон | 5800 |" in md

    def test_pipe_in_cell_is_escaped(self):
        data = [["Условие", "Значение"], ["A|B", "1"]]
        md = docx_to_markdown(make_docx(tables=[data]))
        assert r"A\|B" in md

    def test_empty_document_returns_empty_string(self):
        md = docx_to_markdown(make_docx())
        assert md == ""

    def test_merged_cells_not_duplicated(self):
        """Horizontally merged cells must appear exactly once, not repeated per column span."""
        doc = Document()
        tbl = doc.add_table(rows=2, cols=3)
        # Merge first two cells of the header row.
        tbl.cell(0, 0).merge(tbl.cell(0, 1))
        tbl.cell(0, 0).text = "Объединено"
        tbl.cell(0, 2).text = "Отдельно"
        tbl.cell(1, 0).text = "A"
        tbl.cell(1, 1).text = "B"
        tbl.cell(1, 2).text = "C"
        buf = BytesIO()
        doc.save(buf)
        md = docx_to_markdown(buf.getvalue())

        header_line = [l for l in md.splitlines() if "Объединено" in l][0]
        # "Объединено" must appear exactly once — not duplicated by the span.
        assert header_line.count("Объединено") == 1


# ===========================================================================
# count_h1_h2 tests
# ===========================================================================


class TestCountH1H2:
    def test_counts_h1_and_h2(self):
        md = "# Title\n\n## Section 1\n\ntext\n\n## Section 2\n\n### Sub"
        assert count_h1_h2(md) == 3

    def test_ignores_h3_and_deeper(self):
        assert count_h1_h2("### Only H3\n\n#### H4") == 0

    def test_empty_string(self):
        assert count_h1_h2("") == 0

    def test_does_not_count_inline_hashes(self):
        # A hash mid-line must not be counted.
        assert count_h1_h2("text # not a heading\n## Real heading") == 1


# ===========================================================================
# XLSX parser tests
# ===========================================================================


class TestXlsxToMarkdownSections:
    def test_single_sheet_basic(self):
        data = make_xlsx({"Лист1": [["Название", "Цена"], ["Бетон", 5800], ["Арматура", 7200]]})
        sections = xlsx_to_markdown_sections(data)

        assert len(sections) == 1
        name, md = sections[0]
        assert name == "Лист1"
        assert "| Название | Цена |" in md
        assert "| Бетон | 5800 |" in md
        assert "| Арматура | 7200 |" in md

    def test_table_has_separator_row(self):
        data = make_xlsx({"S": [["Col1", "Col2"]]})
        _, md = xlsx_to_markdown_sections(data)[0]
        assert "---" in md

    def test_multiple_sheets_returned_in_order(self):
        data = make_xlsx(
            {
                "Материалы": [["Наименование"], ["Бетон"]],
                "Итого": [["Сумма"], [1_000_000]],
            }
        )
        sections = xlsx_to_markdown_sections(data)
        assert len(sections) == 2
        assert sections[0][0] == "Материалы"
        assert sections[1][0] == "Итого"

    def test_empty_sheet_is_skipped(self):
        data = make_xlsx(
            {
                "Данные": [["A", "B"], ["1", "2"]],
                "Пустой": [],
            }
        )
        sections = xlsx_to_markdown_sections(data)
        assert len(sections) == 1
        assert sections[0][0] == "Данные"

    def test_all_none_rows_are_skipped(self):
        data = make_xlsx({"S": [["A", "B"], [None, None], ["1", "2"]]})
        _, md = xlsx_to_markdown_sections(data)[0]
        lines = md.strip().split("\n")
        # header + separator + 1 data row (empty row skipped)
        assert len(lines) == 3

    def test_whole_number_float_formatted_without_dot(self):
        # openpyxl may read integer cells as floats.
        data = make_xlsx({"S": [["Value"], [100.0]]})
        _, md = xlsx_to_markdown_sections(data)[0]
        assert "| 100 |" in md

    def test_fractional_float_kept(self):
        data = make_xlsx({"S": [["Value"], [120.5]]})
        _, md = xlsx_to_markdown_sections(data)[0]
        assert "| 120.5 |" in md

    def test_none_cell_rendered_as_empty(self):
        data = make_xlsx({"S": [["A", "B"], ["x", None]]})
        _, md = xlsx_to_markdown_sections(data)[0]
        assert "| x |" in md  # trailing empty cell may be stripped

    def test_rows_with_varying_widths_padded(self):
        # Second row has fewer columns than the header.
        data = make_xlsx({"S": [["A", "B", "C"], ["1"]]})
        _, md = xlsx_to_markdown_sections(data)[0]
        # Data row must have the same column count as the header.
        data_row = [line for line in md.split("\n") if line.startswith("| 1")][0]
        assert data_row.count("|") == 4  # 3 cols → 4 pipes

    def test_date_cell_formatted_as_iso_date(self):
        """openpyxl returns datetime objects for date cells; must render as YYYY-MM-DD."""
        from datetime import datetime

        wb = Workbook()
        wb.remove(wb.active)
        ws = wb.create_sheet("Сроки")
        ws.append(["Дата начала", "Статус"])
        ws.append([datetime(2024, 6, 15), "Активен"])
        buf = BytesIO()
        wb.save(buf)

        _, md = xlsx_to_markdown_sections(buf.getvalue())[0]
        # Must be human-readable date, NOT "2024-06-15 00:00:00".
        assert "2024-06-15" in md
        assert "00:00:00" not in md
