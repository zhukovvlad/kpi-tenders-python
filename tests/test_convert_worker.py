"""Unit tests for the convert worker's _handle() dispatcher.

These tests exercise the full _handle() function (docx/xlsx dispatch +
result_payload shape) without touching the Celery task machinery or any
external service.  A MagicMock replaces MinIOClient so no real MinIO is needed.
"""

from io import BytesIO
from unittest.mock import MagicMock

import pytest
from docx import Document
from openpyxl import Workbook

from app.storage.minio_client import MinIOClient
from app.workers.convert import _handle, _md_object_name

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_minio(return_path: str = "tenders/result.md") -> MagicMock:
    """Return a mock MinIOClient with default_bucket set."""
    mock = MagicMock(spec=MinIOClient)
    mock.default_bucket = "tenders"
    mock.upload.return_value = return_path
    return mock


def _uploaded_content(mock: MagicMock) -> str:
    """Decode the bytes that were passed as second arg to minio.upload()."""
    return mock.upload.call_args[0][1].decode("utf-8")


def _docx(
    headings: list[tuple[int, str]] | None = None,
    paragraphs: list[str] | None = None,
) -> bytes:
    doc = Document()
    for level, text in headings or []:
        doc.add_heading(text, level=level)
    for text in paragraphs or []:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xlsx(sheets: dict[str, list[list]]) -> bytes:
    wb = Workbook()
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(name)
        for row in rows:
            ws.append(row)
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ===========================================================================
# _md_object_name helper
# ===========================================================================


class TestMdObjectName:
    def test_strips_bucket_prefix(self):
        assert _md_object_name("tenders/docs/uuid.docx", "tenders") == "docs/uuid.md"

    def test_non_bucket_prefix_kept(self):
        assert _md_object_name("other/docs/uuid.docx", "tenders") == "other/docs/uuid.md"

    def test_nested_path(self):
        assert _md_object_name("tenders/a/b/c.xlsx", "tenders") == "a/b/c.md"


# ===========================================================================
# DOCX dispatch
# ===========================================================================


class TestHandleDocx:
    def test_format_field(self):
        minio = _make_minio()
        result = _handle(_docx(headings=[(1, "Doc")]), "tenders/file.docx", minio)
        assert result["format"] == "markdown"

    def test_md_storage_path_from_minio(self):
        minio = _make_minio("tenders/file.md")
        result = _handle(_docx(paragraphs=["x"]), "tenders/file.docx", minio)
        assert result["md_storage_path"] == "tenders/file.md"

    def test_no_content_field(self):
        minio = _make_minio()
        result = _handle(_docx(paragraphs=["x"]), "tenders/file.docx", minio)
        assert "content" not in result

    def test_upload_called_once(self):
        minio = _make_minio()
        _handle(_docx(paragraphs=["x"]), "tenders/file.docx", minio)
        minio.upload.assert_called_once()

    def test_uploaded_content_contains_heading(self):
        minio = _make_minio()
        _handle(_docx(headings=[(1, "Договор подряда")]), "tenders/contract.docx", minio)
        assert "# Договор подряда" in _uploaded_content(minio)

    def test_char_count_matches_uploaded_content(self):
        minio = _make_minio()
        result = _handle(_docx(paragraphs=["Текст"]), "tenders/doc.docx", minio)
        assert result["char_count"] == len(_uploaded_content(minio))

    def test_section_count_h1_and_h2_only(self):
        minio = _make_minio()
        result = _handle(
            _docx(headings=[(1, "H1"), (2, "H2"), (3, "H3")]),
            "tenders/file.docx",
            minio,
        )
        # H3 must NOT be counted.
        assert result["section_count"] == 2

    def test_no_sheet_count_for_docx(self):
        minio = _make_minio()
        result = _handle(_docx(paragraphs=["text"]), "tenders/a.docx", minio)
        assert "sheet_count" not in result

    def test_path_casing_does_not_matter(self):
        minio = _make_minio()
        result = _handle(_docx(paragraphs=["x"]), "DIR/file.DOCX", minio)
        assert result["format"] == "markdown"


# ===========================================================================
# XLSX dispatch
# ===========================================================================


class TestHandleXlsx:
    def test_format_field(self):
        minio = _make_minio()
        result = _handle(
            _xlsx({"Лист1": [["A"], ["1"]]}),
            "tenders/report.xlsx",
            minio,
        )
        assert result["format"] == "markdown"

    def test_md_storage_path_from_minio(self):
        minio = _make_minio("tenders/report.md")
        result = _handle(_xlsx({"S": [["C"], [1]]}), "tenders/report.xlsx", minio)
        assert result["md_storage_path"] == "tenders/report.md"

    def test_no_content_field(self):
        minio = _make_minio()
        result = _handle(_xlsx({"S": [["C"], [1]]}), "tenders/f.xlsx", minio)
        assert "content" not in result

    def test_sheet_count(self):
        minio = _make_minio()
        result = _handle(
            _xlsx({"L1": [["A"], [1]], "L2": [["B"], [2]]}),
            "tenders/multi.xlsx",
            minio,
        )
        assert result["sheet_count"] == 2

    def test_no_section_count_for_xlsx(self):
        """section_count was removed from the xlsx payload to avoid confusion with sheet_count."""
        minio = _make_minio()
        result = _handle(_xlsx({"S": [["X"], [1]]}), "tenders/f.xlsx", minio)
        assert "section_count" not in result

    def test_uploaded_content_has_sheet_headers(self):
        minio = _make_minio()
        _handle(
            _xlsx({"Материалы": [["A"], [1]], "Итого": [["B"], [2]]}),
            "tenders/budget.xlsx",
            minio,
        )
        uploaded = _uploaded_content(minio)
        assert "## Лист: Материалы" in uploaded
        assert "## Лист: Итого" in uploaded

    def test_char_count_matches_uploaded_content(self):
        minio = _make_minio()
        result = _handle(_xlsx({"S": [["Col"], ["val"]]}), "tenders/f.xlsx", minio)
        assert result["char_count"] == len(_uploaded_content(minio))


# ===========================================================================
# Error cases
# ===========================================================================


class TestHandleErrors:
    def test_unsupported_extension_raises_value_error(self):
        with pytest.raises(ValueError, match="Unsupported"):
            _handle(b"data", "document.pdf", _make_minio())

    def test_no_extension_raises_value_error(self):
        with pytest.raises(ValueError, match="Unsupported"):
            _handle(b"data", "no_extension", _make_minio())

    def test_empty_docx_raises_value_error(self):
        """A DOCX with no content must raise ValueError, not upload an empty file."""
        with pytest.raises(ValueError, match="empty DOCX"):
            _handle(_docx(), "tenders/file.docx", _make_minio())

    def test_empty_xlsx_raises_value_error(self):
        """An XLSX where all sheets are empty must raise ValueError."""
        empty = _xlsx({"Пустой": []})
        with pytest.raises(ValueError, match="empty XLSX"):
            _handle(empty, "tenders/file.xlsx", _make_minio())

    def test_minio_upload_error_propagates_docx(self):
        """MinIOError from minio.upload() must propagate out of _handle_docx."""
        from app.storage.minio_client import MinIOError

        minio = _make_minio()
        minio.upload.side_effect = MinIOError("S3 error")
        with pytest.raises(MinIOError):
            _handle(_docx(paragraphs=["text"]), "tenders/f.docx", minio)

    def test_minio_upload_error_propagates_xlsx(self):
        """MinIOError from minio.upload() must propagate out of _handle_xlsx."""
        from app.storage.minio_client import MinIOError

        minio = _make_minio()
        minio.upload.side_effect = MinIOError("S3 error")
        with pytest.raises(MinIOError):
            _handle(_xlsx({"S": [["Col"], ["val"]]}), "tenders/f.xlsx", minio)
